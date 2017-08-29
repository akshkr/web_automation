import getc
import docx
import getpic
import goog

"""This is the main driver file of the application Website automation. Run this file with Python 3 along with the packages required and selenium."""

print("\n\nHI Akashh! We are preparing the data for your website. Please wait!")
doc = docx.Document()

"""The following code helps to traverse through a partiular list of the given websites. This technology is outdated according to the current technology of the application. These lines of codes uses the getc.py file."""

"""s= ['http://www.digit.in/', 'https://www.cnet.com/', 'https://techcrunch.com/', 'https://thenextweb.com/', 'http://tech.firstpost.com/']
#s = ['']
r = 'nokia'
url_number = 0
for url in s:
    url_number = url_number + 1
    t = getc.get_content(url, r)
#    print(t)
#    print("\n\n")"""


"""The following lines are the main running part of the program. A docx file is initialised at the top of the program.Following lines asks for user input of the topic to be searched for and on which the articles and the photos will have to be prepared"""

r = str(input("Enter the topic"))

"""The following lines are used to get the data from goog.py which googles the topic, traverses through the results and stores the required text in the variable 't'. 't' has a number of paragraphs appended to it."""

t = goog.get_data(r)

"""The following lines are used to format the text into the docx file and save with the name 'asdfg.docx'"""

url_number = 0
for txt in t:
    url_number = url_number + 1
    doc.add_paragraph("------" + str(url_number) + "------")
    doc.add_paragraph(txt)
doc.save("asdfg.docx")

print("\n.\n.\nHello akash! The text file is prepared.")
print("\nThe pics will be ready shortly.")

"""The following lines are used to save the images obtained from google using selenium."""

getpic.mai(r, 50)
print("\n.\n.\nThe pics are ready!!")
