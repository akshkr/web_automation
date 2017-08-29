import bs4 as bs
import urllib.request
import docx
from urllib.parse import *
from google import search

doc = docx.Document()

"""The following fucntion extractes the body of the contents of the pages in the result of the google search of the topic given. It tries out the list of urls in the result page. Some pages do not allow requests to be made and this raises an exception. We just simply pass any kind of exceptions raised while making the request to make the application stable."""

"""The application stores only those paragraphs whose number of letters exceed 200"""

def get_data(content):
	#ip=input("What would you like to search for? ")
	tex = []
	for url in search(content, stop=20):
		try:
			present_page = urllib.request.urlopen(url).read()
		except Exception as e:
			pass
		present_soup = bs.BeautifulSoup(present_page, 'lxml')
		present_body = present_soup.body
		for para in present_body.find_all('p'):
			q = para.text
			if len(q) > 200:
				tex.append(para.text)
	return tex
"""
for txt in tex:
	doc.add_paragraph("------" + "------")
	doc.add_paragraph(txt)
doc.save("asdfg.docx")"""
